# 检查字体
from docx import Document
from docx.oxml.ns import qn

from base.get_paths import get_paths
from base.get_file_info import get_file_info
from base.write_to_txt import write_to_txt

def check_text(file_path):
    """
    检查 Word 文档中使用的字体，并返回一个字体名称的集合。

    :param doc_path: Word 文档的完整路径
    :return: 使用的字体名称集合
    """

    # 打开 Word 文档
    doc = Document(file_path)

    # 用于存储字体的集合
    fonts = set()

    # 遍历文档中的段落
    for paragraph in doc.paragraphs:
        # 遍历段落中的文本范围
        for run in paragraph.runs:
            get_used_fonts(run, fonts)
    
    # 遍历文档中的表格
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for run in cell.runs:
                    get_used_fonts(run, fonts)
    
    # 遍历文档中的页眉和页脚
    for header_footer in doc.sections:
        for paragraph in header_footer.header.paragraphs:
            for run in paragraph.runs:
                get_used_fonts(run, fonts)
        for paragraph in header_footer.footer.paragraphs:
            for run in paragraph.runs:
                get_used_fonts(run, fonts)
    
    # 遍历文档中的脚注和尾注
    for footnote in doc.footnotes:
        for run in footnote.paragraph.runs:
            get_used_fonts(run, fonts)
    for endnote in doc.endnotes:
        for run in endnote.paragraph.runs:
            get_used_fonts(run, fonts)
    
    # 遍历文本框
    for textbox in doc.inline_shapes:
        for paragraph in textbox.text_frame.paragraphs:
            for run in paragraph.runs:
                get_used_fonts(run, fonts)
    

    return fonts

def get_used_fonts(run, fonts):
    # 添加字体名称到集合中
    try:
        fonts.add(run.font.name)
        fonts.add(run._element.rPr.rFonts.get(qn('w:ascii')))
        fonts.add(run._element.rPr.rFonts.get(qn('w:eastAsia')))
    except Exception as e:
        print(f"Error: {run.text} - {e}")

def process_fonts(callback):
    file_filter = [('Word 文档', '*.docx')]
    file_paths = get_paths('打开文件', file_filter=file_filter)
    print(file_paths)

    # 检查是否选择了文件
    if file_paths is None:
        callable("已取消")
        return None

    results = []
    work_directory = get_file_info(file_paths[0], 'directory')
    
    # 窗口进度标签信息
    total_files = len(file_paths)
    current_file = 0
    callback(f"进度：{current_file} / {total_files}，请不要关闭窗口")

    for file_path in file_paths:
        filename = get_file_info(file_path, 'all_name')
        results.append(filename)

        fonts = get_used_fonts(file_path)
        results.append(f"使用的字体: {fonts}")

        current_file += 1
        callback(f"进度：{current_file} / {total_files}，请不要关闭窗口")
    
    write_to_txt(results, work_directory, '000_检查字体.txt')

    callback(f"已处理 {total_files} 个文件，报告保存在文件目录下")
    print("检查字体已完成")