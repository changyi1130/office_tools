# 检查字体
from docx import Document
from docx.oxml.ns import qn

from file_processing.open_multiple_files import open_multiple_files
from other_functions.extract_file_name import extract_file_name
from other_functions.write_text import write_text

def check_text(file_path, function):
    """
    遍历文档中所有 run

    :param doc_path: Word 文档的完整路径
    :return: 使用的字体名称集合
    """

    # 打开 Word 文档
    doc = Document(file_path)

    # 检查信息的集合
    information = set()

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            get_used_fonts(run, information)

    for table in doc.tables:
        for cell in table._cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    get_used_fonts(run, information)

    for header_footer in doc.sections:
        for paragraph in header_footer.header.paragraphs:
            for run in paragraph.runs:
                get_used_fonts(run, information)
        for paragraph in header_footer.footer.paragraphs:
            for run in paragraph.runs:
                get_used_fonts(run, information)
    
    information.discard(None)
    return information

def get_used_fonts(run, fonts):
    """添加字体名称到集合中"""
    try:
        fonts.add(run.font.name)
        fonts.add(run._element.rPr.rFonts.get(qn('w:ascii')))
        fonts.add(run._element.rPr.rFonts.get(qn('w:eastAsia')))

        run.font.name

    except Exception as e:
        print(f"错误: {run.text} - {e}")

def process_check_text(update_info):
    """选择文档，检查字体"""
    file_filter = [('Word 文档', '*.docx')]
    file_paths = open_multiple_files('打开文件', file_filter=file_filter)

    # 检查是否选择了文件
    if file_paths is None:
        update_info("已取消")
        return None

    results = []
    work_directory = extract_file_name(file_paths[0], 'directory')
    
    # 更新提示信息
    total_files = len(file_paths)
    current_file = 0
    update_info(f"进度：{current_file} / {total_files}，请稍后……")

    for file_path in file_paths:
        filename = extract_file_name(file_path, 'full_name')
        results.append(filename)

        fonts = check_text(file_path, get_used_fonts)
        results.append(f"使用的字体: {fonts}")

        current_file += 1
        update_info(f"进度：{current_file} / {total_files}，请稍后……")
    
    write_text(results, work_directory, '000_检查字体.txt')

    update_info(f"已检查 {total_files} 个文件，报告保存在文件目录下")

    print("检查字体完成")